"""
ContracAI - 文档解析器

合同文档的混合解析策略：
  - 文本型 PDF -> pdfplumber 文本提取 -> LLM 结构化解析
  - 扫描件 PDF / 图像 -> 转换为图像 -> LLM 多模态视觉解析
  - Word .docx -> python-docx 文本提取 -> LLM 结构化解析
"""

import io
import json
import re
import tempfile
import asyncio
import httpx
import uuid
import zipfile
from pathlib import Path
from typing import Any

from ...core.llm import LLMClient
from ...core.config import get_config


EXTRACTION_PROMPT = """你是一个专业的合同信息提取助手。请从以下合同文本中仔细提取关键信息，以 JSON 格式返回。

对于每个字段，请同时给出置信度（confidence）：
- "high": 在原文中有明确的、完整的信息
- "medium": 在原文中有部分信息，需要推断
- "low": 在原文中未找到，仅凭经验猜测或无法确定

返回格式如下（注意每个字段都是一个对象，包含 value 和 confidence）：
{{
  "合同名称": {{"value": "xxx项目采购合同", "confidence": "high"}},
  "合同编号": {{"value": "HT-2025-001", "confidence": "high"}},
  "合同类型": {{"value": "采购", "confidence": "medium"}},
  "对方单位名称": {{"value": "某某科技有限公司", "confidence": "high"}},
  "合同金额": {{"value": 500000, "confidence": "high"}},
  "税率": {{"value": 0.13, "confidence": "medium"}},
  "签订时间": {{"value": "2025-03-15", "confidence": "high"}},
  "生效日期": {{"value": "2025-03-15", "confidence": "high"}},
  "截止日期": {{"value": "2026-03-14", "confidence": "high"}},
  "合同支付条款": {{"value": "验收合格后30日内支付全款", "confidence": "high"}},
  "履约保证金": {{"value": "合同金额的5%", "confidence": "medium"}},
  "经办人": {{"value": "张三", "confidence": "low"}},
  "采购方式": {{"value": "公开招标", "confidence": "medium"}},
  "主办部门": {{"value": "信息技术部", "confidence": "low"}},
  "付款时间节点": {{"value": [
    {{"期次": "第1期", "金额": 150000, "时间": "2025-06-01", "条件": "到货验收后"}},
    {{"期次": "第2期", "金额": 350000, "时间": "2025-12-01", "条件": "终验合格后"}}
  ], "confidence": "high"}}
}}

注意：
- 金额统一转换为元（如"50万"转为500000，"伍拾万元"转为500000）
- 税率转换为小数（如"13%"转为0.13）
- 日期统一使用 YYYY-MM-DD 格式
- 如果某个字段在文本中完全未提及，value 设为空字符串（数值设为0），confidence 设为 "low"
- 付款时间节点：如果合同中有分期付款条款，请逐期提取

合同文本内容：
{text}

请仅返回 JSON，不要有其他内容。"""

VISION_EXTRACTION_PROMPT = """你是一个专业的合同信息提取助手。请仔细阅读这份合同文件的图片，提取以下关键信息并以 JSON 格式返回。

对于每个字段，请同时给出置信度（confidence）：
- "high": 在图片中清晰可读
- "medium": 在图片中模糊或部分可读
- "low": 在图片中无法识别

返回格式（每个字段包含 value 和 confidence）：
{
  "合同名称": {"value": "合同/项目名称", "confidence": "high"},
  "合同编号": {"value": "合同编号", "confidence": "high"},
  "合同类型": {"value": "采购/服务/工程等", "confidence": "medium"},
  "对方单位名称": {"value": "供应商/乙方名称", "confidence": "high"},
  "合同金额": {"value": 0, "confidence": "high"},
  "税率": {"value": 0, "confidence": "medium"},
  "签订时间": {"value": "YYYY-MM-DD", "confidence": "high"},
  "生效日期": {"value": "YYYY-MM-DD", "confidence": "high"},
  "截止日期": {"value": "YYYY-MM-DD", "confidence": "high"},
  "合同支付条款": {"value": "付款条件描述", "confidence": "high"},
  "履约保证金": {"value": "保证金情况", "confidence": "medium"},
  "经办人": {"value": "", "confidence": "low"},
  "采购方式": {"value": "", "confidence": "low"},
  "主办部门": {"value": "", "confidence": "low"},
  "付款时间节点": {"value": [], "confidence": "low"}
}

注意：
- 金额统一转换为元（如"50万"转为500000）
- 税率转换为小数（如"13%"转为0.13）
- 日期统一使用 YYYY-MM-DD 格式
- 如果某个字段无法识别，value 设为空字符串，confidence 设为 "low"

请仅返回 JSON，不要有其他内容。"""


class DocParser:
    """将合同文档（PDF、Word、图像）解析为结构化数据。"""

    def __init__(self):
        self.client = LLMClient(get_config().chat_llm)

    async def parse_file(self, file_path: str | Path) -> dict[str, Any]:
        """自动检测文件类型并进行相应解析。"""
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()

        if suffix in (".pdf", ".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tiff"):
            return await self._parse_with_ocr(file_path)
        elif suffix in (".docx", ".doc"):
            return await self._parse_word(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {suffix}")

    async def _parse_with_ocr(self, file_path: Path) -> dict[str, Any]:
        """使用 MinerU 或 PaddleOCR 解析 PDF 或图像，然后提取结构化数据"""
        mineru_key = get_config().mineru_api_key
        md_text = ""
        raw_bboxes = []
        
        if mineru_key:
            try:
                print(f"尝试使用 MinerU 解析: {file_path}")
                md_text, raw_bboxes = await self._extract_with_mineru_async(file_path, mineru_key)
            except Exception as e:
                print(f"MinerU 解析失败 ({e})，降级使用 PaddleOCR")
                md_text, raw_bboxes = await self._extract_with_paddleocr_async(file_path)
        else:
            md_text, raw_bboxes = await self._extract_with_paddleocr_async(file_path)
            
        if not md_text or not md_text.strip():
            raise Exception("OCR 提取结果为空")
        return await self._extract_from_text(md_text, str(file_path), raw_bboxes)

    async def _extract_with_mineru_async(self, file_path: Path, token: str) -> str:
        """使用 MinerU 精准解析 API 提取 Markdown 文本。"""
        file_name = file_path.name
        data_id = str(uuid.uuid4())
        
        batch_url = "https://mineru.net/api/v4/file-urls/batch"
        headers = {
            "Content-Type": "application/json",
            "Authorization": f"Bearer {token}"
        }
        data = {
            "files": [
                {"name": file_name, "data_id": data_id}
            ],
            "model_version": "vlm"
        }
        
        async with httpx.AsyncClient() as client:
            # 1. 申请上传链接
            res = await client.post(batch_url, headers=headers, json=data, timeout=30.0)
            res.raise_for_status()
            result = res.json()
            if result.get("code") != 0:
                raise Exception(f"MinerU 申请上传链接失败: {result.get('msg')}")
                
            batch_id = result["data"]["batch_id"]
            upload_url = result["data"]["file_urls"][0]
            
            # 2. 上传文件
            print(f"MinerU 任务已创建，batch_id: {batch_id}，正在上传文件...")
            import requests
            with open(file_path, "rb") as f:
                file_bytes = f.read()
                # 使用 requests 进行同步上传，避免 httpx 的 chunked encoding 导致 S3 报错
                # 此处必须不要加 Content-Type，否则 OSS 签名可能校验失败返回 403
                put_res = requests.put(upload_url, data=file_bytes, timeout=120.0)
                if put_res.status_code != 200:
                    raise Exception(f"MinerU 上传失败: HTTP {put_res.status_code} - {put_res.text}")
                
            # 3. 轮询结果
            print("MinerU 文件上传成功，等待解析...")
            poll_url = f"https://mineru.net/api/v4/extract-results/batch/{batch_id}"
            poll_headers = {
                "Accept": "*/*",
                "Authorization": f"Bearer {token}"
            }
            
            while True:
                poll_res = await client.get(poll_url, headers=poll_headers, timeout=30.0)
                poll_res.raise_for_status()
                poll_result = poll_res.json()
                
                if poll_result.get("code") != 0:
                    raise Exception(f"MinerU 查询状态失败: {poll_result.get('msg')}")
                    
                extract_result = poll_result["data"]["extract_result"][0]
                state = extract_result["state"]
                
                if state == "done":
                    zip_url = extract_result["full_zip_url"]
                    break
                elif state == "failed":
                    err_msg = extract_result.get("err_msg", "未知错误")
                    raise Exception(f"MinerU 解析失败: {err_msg}")
                    
                print(f"MinerU 任务进行中，当前状态: {state}")
                await asyncio.sleep(3)
                
            # 4. 下载 zip 包并提取 markdown
            print("MinerU 解析完成，正在下载并提取结果...")
            zip_res = await client.get(zip_url, timeout=60.0)
            zip_res.raise_for_status()
            
            with zipfile.ZipFile(io.BytesIO(zip_res.content)) as z:
                print(f"[DEBUG] zip namelist: {z.namelist()}")
                # 寻找 full.md 或其他 md 文件
                md_files = [name for name in z.namelist() if name.endswith('.md')]
                if not md_files:
                    raise Exception("MinerU 解析结果中未找到 Markdown 文件")
                    
                # 默认读取 full.md，如果没有则读第一个
                target_md = "full.md" if "full.md" in md_files else md_files[0]
                md_bytes = z.read(target_md)
                md_text = md_bytes.decode('utf-8')
                
                # 寻找 json 文件并提取 bbox
                raw_bboxes = []
                # 首先读取 layout.json 获取每页宽高
                page_dims = {}
                layout_files = [name for name in z.namelist() if name.endswith('layout.json')]
                if layout_files:
                    try:
                        layout_data = json.loads(z.read(layout_files[0]))
                        if "pdf_info" in layout_data:
                            for page in layout_data["pdf_info"]:
                                p_idx = page.get("page_idx", 0)
                                p_info = page.get("page_info", page)
                                p_w = p_info.get("width", page.get("width", 1))
                                p_h = p_info.get("height", page.get("height", 1))
                                page_dims[p_idx] = (p_w, p_h)
                    except Exception as e:
                        print(f"读取 layout.json 失败: {e}")
                        
                # 兼容性 Fallback：如果 page_dims 为空或全为1，使用 pdfplumber 获取真实的 PDF 宽高（单位: 磅/Points）
                try:
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        for i, page in enumerate(pdf.pages):
                            if i not in page_dims or page_dims[i] == (1, 1):
                                page_dims[i] = (float(page.width), float(page.height))
                except Exception as e:
                    print(f"pdfplumber 回退获取宽高失败: {e}")

                json_files = [name for name in z.namelist() if name.endswith('_content_list.json')]
                if json_files:
                    try:
                        json_data = json.loads(z.read(json_files[0]))
                        if isinstance(json_data, list):
                            for block in json_data:
                                if "text" in block and "bbox" in block:
                                    bbox = block["bbox"]
                                    if len(bbox) == 4:
                                        p_idx = block.get("page_idx", 0)
                                        pw, ph = page_dims.get(p_idx, (1, 1))
                                        
                                        # 如果获取不到真实宽高（还是 1），避免将坐标误认为极大比例
                                        if pw == 1 or ph == 1:
                                            nx0, ny0, nx1, ny1 = bbox[0], bbox[1], bbox[2], bbox[3]
                                        else:
                                            # Normalize coordinates to [0, 1] percentages
                                            nx0 = bbox[0] / pw
                                            ny0 = bbox[1] / ph
                                            nx1 = bbox[2] / pw
                                            ny1 = bbox[3] / ph
                                        
                                        raw_bboxes.append({
                                            "text": block["text"],
                                            "bbox": [nx0, ny0, nx1, ny1, p_idx + 1]
                                        })
                    except Exception as e:
                        print(f"提取 bbox 失败: {e}")
                        
                # 极端情况兜底：如果有些 bbox 还是绝对坐标（比如 pdfplumber 失败且 MinerU 返回绝对像素），则利用页面最大值进行自适应归一化
                page_maxes = {}
                for item in raw_bboxes:
                    box = item["bbox"]
                    p = box[4]
                    if box[0] > 2 or box[1] > 2:
                        mx, my = page_maxes.get(p, (1, 1))
                        page_maxes[p] = (max(mx, box[2]), max(my, box[3]))
                        
                if page_maxes:
                    for p in page_maxes:
                        page_maxes[p] = (page_maxes[p][0] * 1.05, page_maxes[p][1] * 1.05)
                    for item in raw_bboxes:
                        box = item["bbox"]
                        p = box[4]
                        if box[0] > 2 or box[1] > 2:
                            pw, ph = page_maxes.get(p, (1, 1))
                            if pw > 1 and ph > 1:
                                item["bbox"] = [box[0] / pw, box[1] / ph, box[2] / pw, box[3] / ph, p]
                
                print(f"[DEBUG] page_dims: {page_dims}")
                print(f"[DEBUG] raw_bboxes count: {len(raw_bboxes)}")
                if raw_bboxes:
                    print(f"[DEBUG] sample bbox: {raw_bboxes[0]}")
                        
                return md_text, raw_bboxes

    async def _extract_with_paddleocr_async(self, file_path: Path) -> str:
        """使用 PaddleOCR 公网 API 提取 Markdown 文本。"""
        JOB_URL = "https://paddleocr.aistudio-app.com/api/v2/ocr/jobs"
        TOKEN = get_config().paddle_ocr_token or "cb6e5988ccc271b1411cb31b423831e53f8dbd8f"
        MODEL = "PaddleOCR-VL-1.6"

        headers = {
            "Authorization": f"bearer {TOKEN}",
        }
        optional_payload = {
            "useDocOrientationClassify": False,
            "useDocUnwarping": False,
            "useChartRecognition": False,
        }

        data = {
            "model": MODEL,
            "optionalPayload": json.dumps(optional_payload)
        }

        async with httpx.AsyncClient() as client:
            with open(file_path, "rb") as f:
                files = {"file": f}
                job_response = await client.post(JOB_URL, headers=headers, data=data, files=files, timeout=60.0)

            job_response.raise_for_status()
            jobId = job_response.json()["data"]["jobId"]
            print(f"PaddleOCR 任务已提交，job id: {jobId}")

            # 轮询获取结果
            jsonl_url = ""
            while True:
                res = await client.get(f"{JOB_URL}/{jobId}", headers=headers, timeout=10.0)
                res.raise_for_status()
                state = res.json()["data"]["state"]
                
                if state == 'done':
                    jsonl_url = res.json()['data']['resultUrl']['jsonUrl']
                    break
                elif state == "failed":
                    error_msg = res.json()['data']['errorMsg']
                    raise Exception(f"PaddleOCR 任务失败：{error_msg}")
                
                print(f"PaddleOCR 任务进行中，当前状态: {state}")
                await asyncio.sleep(5)

            # 获取解析的 JSONL 并组装成 Markdown 和坐标
            raw_bboxes = []
            if jsonl_url:
                jsonl_res = await client.get(jsonl_url, timeout=30.0)
                jsonl_res.raise_for_status()
                lines = jsonl_res.text.strip().split('\n')
                
                md_text = ""
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    parsed_line = json.loads(line)
                    result = parsed_line.get("result", {})
                    # PaddleOCR 默认是没有页码的，如果有多页会有单独的 jsonl 或者在别的地方
                    # 假定都在一页，或者顺次排布
                    page_num = result.get("page_idx", 1)
                    
                    for res in result.get("layoutParsingResults", []):
                        block_text = res.get("markdown", {}).get("text", "")
                        md_text += block_text + "\n\n"
                        
                        box = res.get("box", res.get("polygon", []))
                        # 如果是多边形 [x1, y1, x2, y2, x3, y3, x4, y4] 转成 [x0, y0, x1, y1]
                        if len(box) >= 8:
                            xs = [box[i] for i in range(0, len(box), 2)]
                            ys = [box[i] for i in range(1, len(box), 2)]
                            bbox = [min(xs), min(ys), max(xs), max(ys), page_num]
                        elif len(box) == 4:
                            bbox = [box[0], box[1], box[2], box[3], page_num]
                        else:
                            bbox = [0, 0, 0, 0, page_num]
                            
                        if block_text.strip():
                            raw_bboxes.append({
                                "text": block_text.strip(),
                                "bbox": bbox
                            })
                            
                # 归一化 PaddleOCR 的绝对坐标为百分比
                page_maxes = {}
                for item in raw_bboxes:
                    box = item["bbox"]
                    p = box[4]
                    mx, my = page_maxes.get(p, (1, 1))
                    page_maxes[p] = (max(mx, box[2]), max(my, box[3]))
                
                # 加上一点边距，避免最大的框刚好碰到边缘
                for p in page_maxes:
                    page_maxes[p] = (page_maxes[p][0] * 1.05, page_maxes[p][1] * 1.05)
                    
                for item in raw_bboxes:
                    box = item["bbox"]
                    p = box[4]
                    pw, ph = page_maxes.get(p, (1, 1))
                    if pw > 1 and ph > 1:
                        item["bbox"] = [
                            box[0] / pw, box[1] / ph,
                            box[2] / pw, box[3] / ph,
                            p
                        ]
                
                print(f"PaddleOCR 解析完成，Markdown 长度: {len(md_text)}")
                return md_text, raw_bboxes
            
            return "", []

    async def _parse_word(self, file_path: Path) -> dict[str, Any]:
        """使用 python-docx 解析 Word 文档。"""
        from docx import Document

        doc = Document(str(file_path))
        text_content = "\n".join(
            para.text for para in doc.paragraphs if para.text.strip()
        )

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                text_content += "\n" + row_text

        return await self._extract_from_text(text_content, str(file_path))

    async def _extract_from_text(
        self, text: str, source: str, raw_bboxes: list = None
    ) -> dict[str, Any]:
        """使用 LLM 从纯文本中提取结构化数据。"""
        prompt = EXTRACTION_PROMPT.format(text=text[:8000])

        response = await self.client.chat(
            messages=[{"role": "user", "content": prompt}]
        )
        content = self.client.get_choice_message(response).get("content", "")
        return self._parse_json_response(content, source, raw_bboxes)

    def _parse_json_response(
        self, response_text: str, source: str, raw_bboxes: list = None
    ) -> dict[str, Any]:
        """解析 LLM 响应中的 JSON，具备多轮容错能力。"""
        text = response_text.strip()

        # 第一轮：移除 markdown 代码块包裹
        if text.startswith("```"):
            lines = text.split("\n")
            lines = [l for l in lines if not l.strip().startswith("```")]
            text = "\n".join(lines)

        # 第二轮：尝试直接解析
        data = self._try_parse_json(text)

        # 第三轮：如果失败，用正则提取 {...} 块
        if data is None:
            json_match = re.search(r'\{[\s\S]*\}', text)
            if json_match:
                data = self._try_parse_json(json_match.group(0))

        # 全部失败，返回原始响应
        if data is None:
            data = {"_raw_response": response_text, "_parse_error": True}

        # 展平 {value, confidence} 结构为前端友好的格式
        data = self._flatten_confidence(data, raw_bboxes, source_file=source)

        data["_source_file"] = source
        return data

    def _try_parse_json(self, text: str) -> dict | None:
        """尝试解析 JSON，包含常见修复（如尾部逗号）。"""
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            pass

        # 修复尾部逗号: ,} 或 ,]
        fixed = re.sub(r',\s*([}\]])', r'\1', text)
        try:
            return json.loads(fixed)
        except json.JSONDecodeError:
            return None

    def _find_bbox_for_value(self, value: Any, raw_bboxes: list) -> list | None:
        if not raw_bboxes or value is None:
            return None
            
        value_str = str(value).strip().lower()
        if not value_str:
            return None
            
        import difflib
        best_bbox = None
        best_ratio = 0.0
        
        for item in raw_bboxes:
            text = str(item.get("text", "")).strip().lower()
            if not text:
                continue
            
            if value_str in text or text in value_str:
                return item.get("bbox")
                
            ratio = difflib.SequenceMatcher(None, value_str, text).ratio()
            if ratio > best_ratio and ratio > 0.6:
                best_ratio = ratio
                best_bbox = item.get("bbox")
                
        return best_bbox

    def _flatten_confidence(self, data: dict, raw_bboxes: list = None, source_file: str = None) -> dict:
        """
        将 LLM 返回的 {value, confidence} 嵌套结构展平。
        优先尝试用 pdfplumber 提取精确边界框，失败则回退到 MinerU/PaddleOCR 的 raw_bboxes。
        """
        confidence_map = {}
        bbox_map = {}
        flat_data = {}
        
        pdf_pages = []
        page_dims = {}
        if source_file and source_file.lower().endswith('.pdf'):
            try:
                import pdfplumber
                # Keep pdf object alive during this function
                pdf = pdfplumber.open(source_file)
                for i, page in enumerate(pdf.pages):
                    page_dims[i+1] = (page.width, page.height)
                    pdf_pages.append((i+1, page))
            except Exception as e:
                print(f"[DEBUG] pdfplumber 加载失败: {e}")

        def find_bbox(value_str):
            value_str = str(value_str).strip()
            if not value_str: return None
            
            # 1. 优先使用 pdfplumber 的全文搜索定位完整短语 (最精准)
            if pdf_pages:
                for p_idx, page in pdf_pages:
                    try:
                        matches = page.search(value_str)
                        if matches:
                            m = matches[0]
                            pw, ph = page_dims.get(p_idx, (595.276, 841.89))
                            return [m["x0"]/pw, m["top"]/ph, m["x1"]/pw, m["bottom"]/ph, p_idx]
                    except:
                        pass
                
                # 1.1 如果没搜到，尝试拆分成单词搜索
                val_lower = value_str.lower()
                for p_idx, page in pdf_pages:
                    try:
                        for w in page.extract_words():
                            text = w["text"].lower()
                            if val_lower in text or text in val_lower:
                                pw, ph = page_dims.get(p_idx, (595.276, 841.89))
                                return [w["x0"]/pw, w["top"]/ph, w["x1"]/pw, w["bottom"]/ph, p_idx]
                    except:
                        pass

            # 2. 退回使用 MinerU 的模糊匹配
            return self._find_bbox_for_value(value_str, raw_bboxes)

        for key, val in data.items():
            if key.startswith("_"):
                flat_data[key] = val
                continue

            if isinstance(val, dict) and "value" in val:
                flat_data[key] = val["value"]
                confidence_map[key] = val.get("confidence", "medium")
                bbox = find_bbox(val["value"])
                if bbox:
                    bbox_map[key] = bbox
            else:
                flat_data[key] = val
                confidence_map[key] = "high"
                bbox = find_bbox(val)
                if bbox:
                    bbox_map[key] = bbox

        flat_data["_confidence"] = confidence_map
        if bbox_map:
            flat_data["_bboxes"] = bbox_map
            
        if raw_bboxes:
            flat_data["_raw_bboxes"] = raw_bboxes
            
        return flat_data


from ...core.config import get_config
cfg = get_config()
parser_instance = DocParser() if cfg.chat_llm and cfg.chat_llm.api_key else None
