"""
ContracAI - 文档解析器

合同文档的混合解析策略：
  - 文本型 PDF -> pdfplumber 文本提取 -> LLM 结构化解析
  - 扫描件 PDF / 图像 -> 转换为图像 -> LLM 多模态视觉解析
  - Word .docx -> python-docx 文本提取 -> LLM 结构化解析
"""

import io
import json
import tempfile
from pathlib import Path
from typing import Any

from .llm_client import LLMClient


EXTRACTION_PROMPT = """请从以下合同文本中提取关键信息，以 JSON 格式返回。如果某个字段在文本中未提及，设为空字符串。

需要提取的字段：
{
  "合同名称": "合同/项目名称",
  "合同编号": "合同编号",
  "合同类型": "采购/服务/工程等",
  "对方单位名称": "供应商/乙方名称",
  "合同金额": 0,
  "税率": 0,
  "签订时间": "YYYY-MM-DD",
  "生效日期": "YYYY-MM-DD",
  "截止日期": "YYYY-MM-DD",
  "合同支付条款": "付款条件描述",
  "履约保证金": "保证金情况",
  "经办人": ""
}

注意：
- 金额统一转换为元（如"50万"转为500000）
- 税率转换为小数（如"13%"转为0.13）
- 日期统一使用 YYYY-MM-DD 格式

合同文本内容：
{text}

请仅返回 JSON，不要有其他内容。
"""

VISION_EXTRACTION_PROMPT = """请仔细阅读这份合同文件的图片，提取以下关键信息并以 JSON 格式返回。如果某个字段无法识别，设为空字符串。

需要提取的字段：
{
  "合同名称": "合同/项目名称",
  "合同编号": "合同编号",
  "合同类型": "采购/服务/工程等",
  "对方单位名称": "供应商/乙方名称",
  "合同金额": 0,
  "税率": 0,
  "签订时间": "YYYY-MM-DD",
  "生效日期": "YYYY-MM-DD",
  "截止日期": "YYYY-MM-DD",
  "合同支付条款": "付款条件描述",
  "履约保证金": "保证金情况",
  "经办人": ""
}

注意：
- 金额统一转换为元（如"50万"转为500000）
- 税率转换为小数（如"13%"转为0.13）
- 日期统一使用 YYYY-MM-DD 格式

请仅返回 JSON，不要有其他内容。
"""


class DocParser:
    """将合同文档（PDF、Word、图像）解析为结构化数据。"""

    def __init__(self):
        self.client = LLMClient()

    async def parse_file(self, file_path: str | Path) -> dict[str, Any]:
        """
        自动检测文件类型并进行相应解析。

        参数:
            file_path: 文档文件路径。

        返回:
            包含提取的合同字段和元数据的字典。
        """
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            return await self._parse_pdf(file_path)
        elif suffix in (".docx", ".doc"):
            return await self._parse_word(file_path)
        elif suffix in (".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tiff"):
            return await self._parse_image(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    async def _parse_pdf(self, file_path: Path) -> dict[str, Any]:
        """解析 PDF - 优先尝试文本提取，失败则回退至视觉 OCR 解析。"""
        import pdfplumber

        text_content = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_content += page_text + "\n"

        # 如果提取到了有意义的文本，使用基于文本的提取
        if len(text_content.strip()) > 100:
            return await self._extract_from_text(text_content, str(file_path))

        # 否则，将页面转换为图片并使用视觉 OCR 解析
        return await self._parse_pdf_as_images(file_path)

    async def _parse_pdf_as_images(self, file_path: Path) -> dict[str, Any]:
        """将 PDF 页面转换为图片，并使用 LLM 视觉进行 OCR 解析。"""
        from pdf2image import convert_from_path

        # 将 PDF 页面转换为图片
        images = convert_from_path(str(file_path), dpi=200)

        # 保存到临时文件
        temp_paths = []
        for i, img in enumerate(images[:5]):  # 限制为前 5 页
            temp_path = Path(tempfile.gettempdir()) / f"contracai_page_{i}.png"
            img.save(str(temp_path), "PNG")
            temp_paths.append(str(temp_path))

        try:
            response_text = await self.client.chat_with_vision(
                VISION_EXTRACTION_PROMPT, temp_paths
            )
            return self._parse_json_response(response_text, str(file_path))
        finally:
            # 清理临时文件
            for p in temp_paths:
                Path(p).unlink(missing_ok=True)

    async def _parse_word(self, file_path: Path) -> dict[str, Any]:
        """使用 python-docx 解析 Word 文档。"""
        from docx import Document

        doc = Document(str(file_path))
        text_content = "\n".join(
            para.text for para in doc.paragraphs if para.text.strip()
        )

        # 同时从表格中提取文本
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                text_content += "\n" + row_text

        return await self._extract_from_text(text_content, str(file_path))

    async def _parse_image(self, file_path: Path) -> dict[str, Any]:
        """使用 LLM 视觉直接解析图像文件。"""
        response_text = await self.client.chat_with_vision(
            VISION_EXTRACTION_PROMPT, [str(file_path)]
        )
        return self._parse_json_response(response_text, str(file_path))

    async def _extract_from_text(
        self, text: str, source: str
    ) -> dict[str, Any]:
        """使用 LLM 从纯文本中提取结构化 data。"""
        prompt = EXTRACTION_PROMPT.format(text=text[:8000])  # 如果太长则截断

        response = await self.client.chat(
            messages=[{"role": "user", "content": prompt}]
        )
        content = self.client.get_choice_message(response).get("content", "")
        return self._parse_json_response(content, source)

    def _parse_json_response(
        self, response_text: str, source: str
    ) -> dict[str, Any]:
        """解析 LLM 响应中的 JSON，处理 Markdown 代码块。"""
        text = response_text.strip()

        # 如果存在，移除 markdown 代码块包裹
        if text.startswith("```"):
            lines = text.split("\n")
            # 移除第一行 (```json) 和最后一行 (```)
            lines = [l for l in lines if not l.strip().startswith("```")]
            text = "\n".join(lines)

        try:
            data = json.loads(text)
        except json.JSONDecodeError:
            data = {"_raw_response": response_text, "_parse_error": True}

        data["_source_file"] = source
        return data
